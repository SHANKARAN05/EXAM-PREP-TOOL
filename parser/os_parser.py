import os
import re
from docx import Document

def detect_os_topic(text):
    """Detect topic from question text"""
    text_lower = text.lower()
    
    if 'process' in text_lower or 'thread' in text_lower:
        return 'Process & Threads'
    elif 'schedul' in text_lower or 'fcfs' in text_lower or 'sjf' in text_lower or 'round robin' in text_lower:
        return 'CPU Scheduling'
    elif 'deadlock' in text_lower or 'banker' in text_lower:
        return 'Deadlock'
    elif 'memory' in text_lower or 'paging' in text_lower or 'segment' in text_lower or 'virtual' in text_lower:
        return 'Memory Management'
    elif 'semaphore' in text_lower or 'mutex' in text_lower or 'synchroniz' in text_lower:
        return 'Synchronization'
    elif 'file' in text_lower or 'disk' in text_lower:
        return 'File Systems'
    elif 'system call' in text_lower or 'kernel' in text_lower:
        return 'System Calls'
    else:
        return 'OS Basics'

def parse_os_file(filepath):
    """Parse Operating_System_Interview_Q.docx"""
    questions = []
    seen_questions = set()  # Track duplicates
    
    try:
        doc = Document(filepath)
        current_question = None
        current_answer = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this is a question (bold and starts with number)
            is_bold = False
            if para.runs:
                is_bold = any(run.bold for run in para.runs)
            
            is_question = is_bold and re.match(r'^\d+\.', text)
            
            if is_question:
                # Save previous question if exists
                if current_question:
                    answer_text = '\n'.join(current_answer).strip()
                    if answer_text:
                        # Check for duplicates using hash
                        q_hash = hash(current_question.lower().strip())
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            questions.append({
                                'question_type': 'theory',
                                'question': current_question,
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': answer_text,
                                'explanation': None,
                                'topic': detect_os_topic(current_question)
                            })
                
                # Start new question
                # Remove number prefix
                current_question = re.sub(r'^\d+\.\s*', '', text)
                current_answer = []
            else:
                # This is part of the answer
                if current_question:
                    current_answer.append(text)
        
        # Save last question
        if current_question:
            answer_text = '\n'.join(current_answer).strip()
            if answer_text:
                # Check for duplicates using hash
                q_hash = hash(current_question.lower().strip())
                if q_hash not in seen_questions:
                    seen_questions.add(q_hash)
                    questions.append({
                        'question_type': 'theory',
                        'question': current_question,
                        'option_a': None,
                        'option_b': None,
                        'option_c': None,
                        'option_d': None,
                        'answer': answer_text,
                        'explanation': None,
                        'topic': detect_os_topic(current_question)
                    })
        
        # Handle tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            
            # Add table content to the last question's answer if exists
            if questions and table_text:
                questions[-1]['answer'] += '\n\n' + '\n'.join(table_text)
    
    except Exception as e:
        print(f"Error parsing {filepath}: {e}")
    
    return questions

def parse_os_files():
    """Parse OS file and return questions"""
    questions = []
    
    os_file = os.path.join('data', 'Operating System Interview Q.docx')
    if os.path.exists(os_file):
        questions = parse_os_file(os_file)
        print(f"Parsed {len(questions)} questions from Operating System Interview Q.docx")
    
    return questions
