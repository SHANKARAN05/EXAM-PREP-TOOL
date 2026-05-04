import re
import os
from docx import Document


def normalize_code_for_hashing(code):
    """Normalize code by removing extra whitespace and comments for duplicate detection"""
    # Remove extra whitespace and normalize
    lines = [line.strip() for line in code.split('\n') if line.strip()]
    # Filter out explanation lines (non-code text)
    code_lines = []
    for line in lines:
        # Keep lines that look like code (contain common code patterns)
        if any(pattern in line for pattern in ['(', ')', '{', '}', ';', '=', 'System.', 'String', 'int ', 'class ']):
            code_lines.append(line)
    return '\n'.join(code_lines)


def parse_java_qa_docx(filepath, seen_questions=None):
    """Parse JAVA Q and A.docx file - handles MCQs and theory questions"""
    questions = []
    if seen_questions is None:
        seen_questions = set()  # Track duplicates
    
    try:
        doc = Document(filepath)
        
        # ========== PHASE 1: Parse MCQs (Questions 1-20) ==========
        current_question = None
        current_options = {}
        current_answer = None
        current_explanation = ""
        question_number = 0
        
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                i += 1
                continue
            
            # Stop at PART2
            if 'PART2' in text or 'PART 2' in text:
                break
            
            # Check for numbered question (MCQ format)
            match = re.match(r'^(\d+)\.\s+(.+)', text)
            if match:
                # Save previous question if exists
                if current_question and current_answer and len(current_options) >= 4:
                    q_hash = hash(current_question)
                    if q_hash not in seen_questions:
                        seen_questions.add(q_hash)
                        questions.append({
                            'question_type': 'mcq',
                            'question': current_question,
                            'option_a': current_options.get('A', ''),
                            'option_b': current_options.get('B', ''),
                            'option_c': current_options.get('C', ''),
                            'option_d': current_options.get('D', ''),
                            'answer': current_answer,
                            'explanation': current_explanation,
                            'topic': 'Core Java'
                        })
                
                question_number = int(match.group(1))
                
                # Skip duplicates (10-17)
                if 10 <= question_number <= 17:
                    i += 1
                    continue
                
                current_question = match.group(2).strip()
                current_options = {}
                current_answer = None
                current_explanation = ""
                
                # Look ahead for options
                i += 1
                options_text = ""
                
                while i < len(doc.paragraphs):
                    next_text = doc.paragraphs[i].text.strip()
                    
                    if next_text.startswith('Answer'):
                        break
                    elif next_text and not re.match(r'^\d+\.', next_text):
                        options_text += " " + next_text
                        i += 1
                    else:
                        break
                
                # Parse options
                opt_pattern = r'([A-D])\)\s*([^A-D]+?)(?=\s+[A-D]\)|$)'
                opt_matches = re.findall(opt_pattern, options_text)
                
                for opt_letter, opt_text in opt_matches:
                    current_options[opt_letter] = opt_text.strip()
                
                continue
            
            # Check for Answer
            if text.startswith('Answer'):
                ans_match = re.search(r'Answer\s*:?\s*([A-E])', text)
                if ans_match:
                    current_answer = ans_match.group(1)
                i += 1
                continue
            
            # Check for Explanation
            if text.startswith('Explanation'):
                expl_match = re.search(r'Explanation\s*:?\s*(.+)', text)
                if expl_match:
                    current_explanation = expl_match.group(1).strip()
                i += 1
                continue
            
            i += 1
        
        # Save last MCQ
        if current_question and current_answer and len(current_options) >= 4:
            q_hash = hash(current_question)
            if q_hash not in seen_questions:
                seen_questions.add(q_hash)
                questions.append({
                    'question_type': 'mcq',
                    'question': current_question,
                    'option_a': current_options.get('A', ''),
                    'option_b': current_options.get('B', ''),
                    'option_c': current_options.get('C', ''),
                    'option_d': current_options.get('D', ''),
                    'answer': current_answer,
                    'explanation': current_explanation,
                    'topic': 'Core Java'
                })
        
        # ========== PHASE 2: Parse PART2 Theory Questions ==========
        in_part2 = False
        in_output_section = False
        in_coding_section = False
        current_question = None
        current_answer_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect section boundaries
            if 'PART2' in text or 'PART 2' in text:
                in_part2 = True
                in_output_section = False
                in_coding_section = False
                continue
            elif '⚡' in text and 'OUTPUT' in text:
                # Save any pending question before entering output section
                if current_question and current_answer_lines:
                    answer_text = '\n'.join(current_answer_lines).strip()
                    if answer_text:
                        q_hash = hash(current_question)
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            questions.append({
                                'question_type': 'theory',
                                'question': current_question,
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': answer_text,
                                'explanation': None,
                                'topic': determine_topic(current_question)
                            })
                
                in_part2 = False
                in_output_section = True
                in_coding_section = False
                current_question = None
                current_answer_lines = []
                continue
            elif '🛠️' in text and 'CODING' in text:
                in_part2 = False
                in_output_section = False
                in_coding_section = True
                current_question = None
                current_answer_lines = []
                continue
            elif 'Part 3' in text or 'PART3' in text or 'Part3' in text:
                in_part2 = False
                in_output_section = False
                in_coding_section = False
                break
            
            # Parse PART2 theory questions
            if in_part2 and not in_output_section and not in_coding_section:
                if re.match(r'✅\s*Q\d+:', text):
                    # Save previous question
                    if current_question and current_answer_lines:
                        answer_text = '\n'.join(current_answer_lines).strip()
                        if answer_text:
                            q_hash = hash(current_question)
                            if q_hash not in seen_questions:
                                seen_questions.add(q_hash)
                                questions.append({
                                    'question_type': 'theory',
                                    'question': current_question,
                                    'option_a': None,
                                    'option_b': None,
                                    'option_c': None,
                                    'option_d': None,
                                    'answer': answer_text,
                                    'explanation': None,
                                    'topic': determine_topic(current_question)
                                })
                    
                    # Extract new question
                    q_match = re.search(r'✅\s*Q\d+:\s*(.+)', text)
                    if q_match:
                        current_question = q_match.group(1).strip()
                        current_answer_lines = []
                elif current_question:
                    # Collect answer lines
                    if text.startswith('Answer:'):
                        answer_part = text.split('Answer:', 1)[1].strip()
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif text.startswith('👉'):
                        answer_part = text.replace('👉', '').strip()
                        answer_part = re.sub(r'^(Answer:|Short punch line:)\s*', '', answer_part)
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif not text.startswith('_') and not text.startswith('SECTION'):
                        current_answer_lines.append(text)
        
        # Save last PART2 question
        if current_question and current_answer_lines and in_part2:
            answer_text = '\n'.join(current_answer_lines).strip()
            if answer_text:
                q_hash = hash(current_question)
                if q_hash not in seen_questions:
                    seen_questions.add(q_hash)
                    questions.append({
                        'question_type': 'theory',
                        'question': current_question,
                        'option_a': None,
                        'option_b': None,
                        'option_c': None,
                        'option_d': None,
                        'answer': answer_text,
                        'explanation': None,
                        'topic': determine_topic(current_question)
                    })
        
        # ========== PHASE 3: Parse OUTPUT QUESTIONS ==========
        in_output_section = False
        current_code_lines = []
        current_answer = ""
        current_reason = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if '⚡' in text and 'OUTPUT' in text:
                in_output_section = True
                continue
            elif '🛠️' in text and 'CODING' in text:
                # Save last output question before entering coding section
                if current_code_lines:
                    code = '\n'.join(current_code_lines)
                    question_text = f"What is the output of the following code?\n\n{code}"
                    # Use normalized code for duplicate detection
                    normalized_code = normalize_code_for_hashing(code)
                    q_hash = hash(normalized_code)
                    
                    if q_hash not in seen_questions:
                        seen_questions.add(q_hash)
                        full_answer = current_answer
                        if current_reason:
                            full_answer += f"\n\nReason: {current_reason}"
                        
                        questions.append({
                            'question_type': 'output',
                            'question': question_text,
                            'option_a': None,
                            'option_b': None,
                            'option_c': None,
                            'option_d': None,
                            'answer': full_answer if full_answer else "Refer study material",
                            'explanation': None,
                            'topic': 'Core Java'
                        })
                
                in_output_section = False
                current_code_lines = []
                current_answer = ""
                current_reason = ""
                continue
            
            if in_output_section:
                if re.match(r'✅\s*Q\d+', text):
                    # Save previous output question
                    if current_code_lines:
                        code = '\n'.join(current_code_lines)
                        question_text = f"What is the output of the following code?\n\n{code}"
                        # Use normalized code for duplicate detection
                        normalized_code = normalize_code_for_hashing(code)
                        q_hash = hash(normalized_code)
                        
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            full_answer = current_answer
                            if current_reason:
                                full_answer += f"\n\nReason: {current_reason}"
                            
                            questions.append({
                                'question_type': 'output',
                                'question': question_text,
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': full_answer if full_answer else "Refer study material",
                                'explanation': None,
                                'topic': 'Core Java'
                            })
                    
                    current_code_lines = []
                    current_answer = ""
                    current_reason = ""
                elif text.startswith('👉') and 'Answer:' in text:
                    current_answer = text.split('Answer:', 1)[1].strip()
                elif text.startswith('Reason:'):
                    current_reason = text.split('Reason:', 1)[1].strip()
                elif text and not text.startswith('_') and not text.startswith('Output'):
                    current_code_lines.append(text)
        
        # ========== PHASE 4: Parse CODING QUESTIONS ==========
        in_coding_section = False
        current_question = None
        current_code_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if '🛠️' in text and 'CODING' in text:
                in_coding_section = True
                continue
            elif 'Part 3' in text or 'PART3' in text:
                # Save last coding question
                if current_question and current_code_lines:
                    code = '\n'.join(current_code_lines)
                    answer = f"Solution:\n\n{code}"
                    q_hash = hash(current_question)
                    
                    if q_hash not in seen_questions:
                        seen_questions.add(q_hash)
                        questions.append({
                            'question_type': 'coding',
                            'question': f"Write a Java program to: {current_question}",
                            'option_a': None,
                            'option_b': None,
                            'option_c': None,
                            'option_d': None,
                            'answer': answer,
                            'explanation': None,
                            'topic': 'Programming'
                        })
                
                in_coding_section = False
                break
            
            if in_coding_section:
                if re.match(r'✅\s*Q\d+:', text):
                    # Save previous coding question
                    if current_question and current_code_lines:
                        code = '\n'.join(current_code_lines)
                        answer = f"Solution:\n\n{code}"
                        q_hash = hash(current_question)
                        
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            questions.append({
                                'question_type': 'coding',
                                'question': f"Write a Java program to: {current_question}",
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': answer,
                                'explanation': None,
                                'topic': 'Programming'
                            })
                    
                    # Extract new question
                    q_match = re.search(r'✅\s*Q\d+:\s*(.+)', text)
                    if q_match:
                        current_question = q_match.group(1).strip()
                        current_code_lines = []
                elif current_question:
                    if 'class' in text.lower() and '{' in text:
                        current_code_lines.append(text)
                    elif current_code_lines and text and not text.startswith('👉') and not text.startswith('Output:'):
                        current_code_lines.append(text)
        
        # ========== PHASE 5: Parse Part 3 Theory Questions ==========
        in_part3 = False
        current_question = None
        current_answer_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            if 'Part 3' in text or 'PART3' in text or 'Part3' in text:
                in_part3 = True
                continue
            elif 'Part 4' in text or 'PART4' in text or 'Part4' in text:
                break
            
            if in_part3:
                # Check for question markers (Q1:, ❓ Q2:, etc.)
                if re.match(r'[❓]?\s*Q\d+:', text):
                    # Save previous question
                    if current_question and current_answer_lines:
                        answer_text = '\n'.join(current_answer_lines).strip()
                        if answer_text:
                            q_hash = hash(current_question)
                            if q_hash not in seen_questions:
                                seen_questions.add(q_hash)
                                questions.append({
                                    'question_type': 'theory',
                                    'question': current_question,
                                    'option_a': None,
                                    'option_b': None,
                                    'option_c': None,
                                    'option_d': None,
                                    'answer': answer_text,
                                    'explanation': None,
                                    'topic': determine_topic(current_question)
                                })
                    
                    # Extract new question
                    q_match = re.search(r'[❓]?\s*Q\d+:\s*(.+)', text)
                    if q_match:
                        current_question = q_match.group(1).strip()
                        current_answer_lines = []
                elif current_question and not text.startswith('SECTION') and not text.startswith('⚙️') and not text.startswith('⚡') and not text.startswith('🧩') and not text.startswith('🔥') and not text.startswith('🛠️'):
                    # Collect answer lines
                    if text.startswith('Answer:'):
                        answer_part = text.split('Answer:', 1)[1].strip()
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif text.startswith('👉'):
                        answer_part = text.replace('👉', '').strip()
                        answer_part = re.sub(r'^Answer:\s*', '', answer_part)
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif not text.startswith('_'):
                        current_answer_lines.append(text)
        
        # Save last Part 3 question
        if current_question and current_answer_lines:
            answer_text = '\n'.join(current_answer_lines).strip()
            if answer_text:
                q_hash = hash(current_question)
                if q_hash not in seen_questions:
                    seen_questions.add(q_hash)
                    questions.append({
                        'question_type': 'theory',
                        'question': current_question,
                        'option_a': None,
                        'option_b': None,
                        'option_c': None,
                        'option_d': None,
                        'answer': answer_text,
                        'explanation': None,
                        'topic': determine_topic(current_question)
                    })
    
    except Exception as e:
        print(f"Error parsing {filepath}: {e}")
        import traceback
        traceback.print_exc()
    
    return questions


def parse_java_q_docx(filepath, seen_questions=None):
    """Parse java Q.docx file"""
    questions = []
    if seen_questions is None:
        seen_questions = set()
    
    try:
        doc = Document(filepath)
        
        # ========== Parse PART2 - Interview Questions ==========
        in_part2 = False
        in_output_section = False
        in_coding_section = False
        current_question = None
        current_answer_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect section boundaries
            if 'PART2' in text or 'PART 2' in text:
                in_part2 = True
                in_output_section = False
                in_coding_section = False
                continue
            elif '⚡' in text and 'OUTPUT' in text:
                # Save pending question
                if current_question and current_answer_lines:
                    answer_text = '\n'.join(current_answer_lines).strip()
                    if answer_text:
                        q_hash = hash(current_question)
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            questions.append({
                                'question_type': 'theory',
                                'question': current_question,
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': answer_text,
                                'explanation': None,
                                'topic': determine_topic(current_question)
                            })
                
                in_part2 = False
                in_output_section = True
                in_coding_section = False
                current_question = None
                current_answer_lines = []
                continue
            elif '🛠️' in text and 'CODING' in text:
                in_part2 = False
                in_output_section = False
                in_coding_section = True
                current_question = None
                current_answer_lines = []
                continue
            elif 'PART3' in text or 'Part 3' in text:
                in_part2 = False
                in_output_section = False
                in_coding_section = False
                break
            
            # Parse PART2 theory questions
            if in_part2:
                if re.match(r'✅\s*Q\d+:', text):
                    # Save previous question
                    if current_question and current_answer_lines:
                        answer_text = '\n'.join(current_answer_lines).strip()
                        if answer_text:
                            q_hash = hash(current_question)
                            if q_hash not in seen_questions:
                                seen_questions.add(q_hash)
                                questions.append({
                                    'question_type': 'theory',
                                    'question': current_question,
                                    'option_a': None,
                                    'option_b': None,
                                    'option_c': None,
                                    'option_d': None,
                                    'answer': answer_text,
                                    'explanation': None,
                                    'topic': determine_topic(current_question)
                                })
                    
                    # Extract new question
                    q_match = re.search(r'✅\s*Q\d+:\s*(.+)', text)
                    if q_match:
                        current_question = q_match.group(1).strip()
                        current_answer_lines = []
                elif current_question:
                    if text.startswith('Answer:'):
                        answer_part = text.split('Answer:', 1)[1].strip()
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif text.startswith('👉'):
                        answer_part = text.replace('👉', '').strip()
                        answer_part = re.sub(r'^(Answer:|Short punch line:)\s*', '', answer_part)
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif not text.startswith('_'):
                        current_answer_lines.append(text)
        
        # ========== Parse OUTPUT QUESTIONS ==========
        in_output_section = False
        current_code_lines = []
        current_answer = ""
        current_reason = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if '⚡' in text and 'OUTPUT' in text:
                in_output_section = True
                continue
            elif '🛠️' in text:
                # Save last output question
                if current_code_lines:
                    code = '\n'.join(current_code_lines)
                    question_text = f"What is the output of the following code?\n\n{code}"
                    # Use normalized code for duplicate detection
                    normalized_code = normalize_code_for_hashing(code)
                    q_hash = hash(normalized_code)
                    
                    if q_hash not in seen_questions:
                        seen_questions.add(q_hash)
                        full_answer = current_answer
                        if current_reason:
                            full_answer += f"\n\nReason: {current_reason}"
                        
                        questions.append({
                            'question_type': 'output',
                            'question': question_text,
                            'option_a': None,
                            'option_b': None,
                            'option_c': None,
                            'option_d': None,
                            'answer': full_answer if full_answer else "Refer study material",
                            'explanation': None,
                            'topic': 'Core Java'
                        })
                
                in_output_section = False
                current_code_lines = []
                current_answer = ""
                current_reason = ""
                continue
            
            if in_output_section:
                if re.match(r'✅\s*Q\d+', text):
                    # Save previous output question
                    if current_code_lines:
                        code = '\n'.join(current_code_lines)
                        question_text = f"What is the output of the following code?\n\n{code}"
                        # Use normalized code for duplicate detection
                        normalized_code = normalize_code_for_hashing(code)
                        q_hash = hash(normalized_code)
                        
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            full_answer = current_answer
                            if current_reason:
                                full_answer += f"\n\nReason: {current_reason}"
                            
                            questions.append({
                                'question_type': 'output',
                                'question': question_text,
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': full_answer if full_answer else "Refer study material",
                                'explanation': None,
                                'topic': 'Core Java'
                            })
                    
                    current_code_lines = []
                    current_answer = ""
                    current_reason = ""
                elif text.startswith('👉') and 'Answer:' in text:
                    current_answer = text.split('Answer:', 1)[1].strip()
                elif text.startswith('Reason:'):
                    current_reason = text.split('Reason:', 1)[1].strip()
                elif text and not text.startswith('_'):
                    current_code_lines.append(text)
        
        # ========== Parse CODING QUESTIONS ==========
        in_coding_section = False
        current_question = None
        current_code_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if '🛠️' in text and 'CODING' in text:
                in_coding_section = True
                continue
            elif 'PART3' in text or 'Part 3' in text:
                # Save last coding question
                if current_question and current_code_lines:
                    code = '\n'.join(current_code_lines)
                    answer = f"Solution:\n\n{code}"
                    q_hash = hash(current_question)
                    
                    if q_hash not in seen_questions:
                        seen_questions.add(q_hash)
                        questions.append({
                            'question_type': 'coding',
                            'question': f"Write a Java program to: {current_question}",
                            'option_a': None,
                            'option_b': None,
                            'option_c': None,
                            'option_d': None,
                            'answer': answer,
                            'explanation': None,
                            'topic': 'Programming'
                        })
                
                in_coding_section = False
                break
            
            if in_coding_section:
                if re.match(r'✅\s*Q\d+:', text):
                    # Save previous coding question
                    if current_question and current_code_lines:
                        code = '\n'.join(current_code_lines)
                        answer = f"Solution:\n\n{code}"
                        q_hash = hash(current_question)
                        
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            questions.append({
                                'question_type': 'coding',
                                'question': f"Write a Java program to: {current_question}",
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': answer,
                                'explanation': None,
                                'topic': 'Programming'
                            })
                    
                    # Extract new question
                    q_match = re.search(r'✅\s*Q\d+:\s*(.+)', text)
                    if q_match:
                        current_question = q_match.group(1).strip()
                        current_code_lines = []
                elif current_question:
                    if 'class' in text.lower() and '{' in text:
                        current_code_lines.append(text)
                    elif current_code_lines and text and not text.startswith('👉') and not text.startswith('Output:'):
                        current_code_lines.append(text)
        
        # ========== Parse PART3 - Additional theory questions ==========
        in_part3 = False
        current_question = None
        current_answer_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            if 'PART3' in text or 'Part 3' in text or 'PART 3' in text:
                in_part3 = True
                continue
            elif 'PART4' in text or 'Part 4' in text:
                break
            
            if in_part3:
                if re.match(r'[❓]?\s*Q\d+:', text):
                    # Save previous question
                    if current_question and current_answer_lines:
                        answer_text = '\n'.join(current_answer_lines).strip()
                        if answer_text:
                            q_hash = hash(current_question)
                            if q_hash not in seen_questions:
                                seen_questions.add(q_hash)
                                questions.append({
                                    'question_type': 'theory',
                                    'question': current_question,
                                    'option_a': None,
                                    'option_b': None,
                                    'option_c': None,
                                    'option_d': None,
                                    'answer': answer_text,
                                    'explanation': None,
                                    'topic': determine_topic(current_question)
                                })
                    
                    # Extract new question
                    q_match = re.search(r'[❓]?\s*Q\d+:\s*(.+)', text)
                    if q_match:
                        current_question = q_match.group(1).strip()
                        current_answer_lines = []
                elif current_question and not text.startswith('SECTION'):
                    if text.startswith('Answer:'):
                        answer_part = text.split('Answer:', 1)[1].strip()
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif text.startswith('👉'):
                        answer_part = text.replace('👉', '').strip()
                        answer_part = re.sub(r'^Answer:\s*', '', answer_part)
                        if answer_part:
                            current_answer_lines.append(answer_part)
                    elif not text.startswith('_') and not text.startswith('⚙️') and not text.startswith('⚡') and not text.startswith('🧩') and not text.startswith('🔥') and not text.startswith('🛠️'):
                        current_answer_lines.append(text)
        
        # Save last PART3 question
        if current_question and current_answer_lines:
            answer_text = '\n'.join(current_answer_lines).strip()
            if answer_text:
                q_hash = hash(current_question)
                if q_hash not in seen_questions:
                    seen_questions.add(q_hash)
                    questions.append({
                        'question_type': 'theory',
                        'question': current_question,
                        'option_a': None,
                        'option_b': None,
                        'option_c': None,
                        'option_d': None,
                        'answer': answer_text,
                        'explanation': None,
                        'topic': determine_topic(current_question)
                    })
        
    except Exception as e:
        print(f"Error parsing {filepath}: {e}")
        import traceback
        traceback.print_exc()
    
    return questions


def determine_topic(question_text):
    """Determine topic from question text"""
    q_lower = question_text.lower()
    
    if 'collection' in q_lower or 'arraylist' in q_lower or 'hashmap' in q_lower or 'list' in q_lower or 'set' in q_lower:
        return 'Collections'
    elif 'thread' in q_lower or 'synchroniz' in q_lower or 'runnable' in q_lower or 'deadlock' in q_lower or 'volatile' in q_lower:
        return 'Multithreading'
    elif 'jvm' in q_lower or 'jre' in q_lower or 'jdk' in q_lower or 'memory' in q_lower or 'garbage' in q_lower or 'heap' in q_lower or 'stack' in q_lower:
        return 'JVM & Memory'
    elif 'exception' in q_lower or 'throw' in q_lower or 'catch' in q_lower or 'finally' in q_lower:
        return 'Exception Handling'
    elif 'string' in q_lower:
        return 'Strings'
    elif 'overload' in q_lower or 'overrid' in q_lower or 'polymorphism' in q_lower or 'inheritance' in q_lower or 'encapsulation' in q_lower:
        return 'OOPS'
    else:
        return 'Core Java'


def parse_java_files():
    """Parse both Java DOCX files and return combined questions"""
    questions = []
    seen_questions = set()  # Global duplicate tracking across both files
    
    # Parse JAVA Q and A.docx FIRST (it has complete answers)
    qa_file = os.path.join('data', 'JAVA Q and A.docx')
    if os.path.exists(qa_file):
        qa_questions = parse_java_qa_docx(qa_file, seen_questions)
        questions.extend(qa_questions)
        print(f"Parsed {len(qa_questions)} questions from JAVA Q and A.docx")
    
    # Parse java Q.docx SECOND (duplicates will be skipped)
    q_file = os.path.join('data', 'java Q.docx')
    if os.path.exists(q_file):
        q_questions = parse_java_q_docx(q_file, seen_questions)
        questions.extend(q_questions)
        print(f"Parsed {len(q_questions)} questions from java Q.docx")
    
    return questions
