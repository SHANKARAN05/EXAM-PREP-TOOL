import os
import re
from docx import Document
from difflib import SequenceMatcher

def similarity(a, b):
    """Calculate similarity between two strings"""
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

def detect_network_topic(text):
    """Detect topic from question text"""
    text_lower = text.lower()
    
    if 'osi' in text_lower or 'tcp/ip' in text_lower or 'layer' in text_lower:
        return 'Network Models'
    elif 'ip' in text_lower or 'ipv4' in text_lower or 'ipv6' in text_lower or 'mac' in text_lower or 'address' in text_lower:
        return 'Addressing'
    elif 'dns' in text_lower or 'dhcp' in text_lower or 'http' in text_lower or 'ftp' in text_lower or 'smtp' in text_lower or 'protocol' in text_lower:
        return 'Protocols'
    elif 'router' in text_lower or 'switch' in text_lower or 'hub' in text_lower or 'firewall' in text_lower or 'gateway' in text_lower:
        return 'Network Devices'
    elif 'vpn' in text_lower or 'encryption' in text_lower or 'security' in text_lower or 'firewall' in text_lower:
        return 'Security'
    else:
        return 'Networking Basics'

def parse_fresher_networking_file(filepath, seen_questions=None):
    """Parse Fresher_Networking_Interview.docx with full Q&A"""
    questions = []
    if seen_questions is None:
        seen_questions = set()
    
    try:
        doc = Document(filepath)
        current_question = None
        current_answer = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this is a question (bold and starts with number or Q)
            is_bold = False
            if para.runs:
                is_bold = any(run.bold for run in para.runs)
            
            is_question = is_bold and (re.match(r'^\d+\.', text) or re.match(r'^Q\d+', text, re.IGNORECASE))
            
            if is_question:
                # Save previous question if exists
                if current_question:
                    answer_text = '\n'.join(current_answer).strip()
                    if answer_text:
                        # Check for duplicates using hash
                        q_hash = hash(current_question.lower().strip())
                        if q_hash not in seen_questions:
                            seen_questions.add(q_hash)
                            questions.append({
                                'question_type': 'theory',
                                'question': current_question,
                                'option_a': None,
                                'option_b': None,
                                'option_c': None,
                                'option_d': None,
                                'answer': answer_text,
                                'explanation': None,
                                'topic': detect_network_topic(current_question)
                            })
                
                # Start new question
                # Remove number prefix
                current_question = re.sub(r'^\d+\.\s*', '', text)
                current_question = re.sub(r'^Q\d+[:\.\s]*', '', current_question, flags=re.IGNORECASE)
                current_answer = []
            else:
                # This is part of the answer
                if current_question:
                    current_answer.append(text)
        
        # Save last question
        if current_question:
            answer_text = '\n'.join(current_answer).strip()
            if answer_text:
                # Check for duplicates using hash
                q_hash = hash(current_question.lower().strip())
                if q_hash not in seen_questions:
                    seen_questions.add(q_hash)
                    questions.append({
                        'question_type': 'theory',
                        'question': current_question,
                        'option_a': None,
                        'option_b': None,
                        'option_c': None,
                        'option_d': None,
                        'answer': answer_text,
                        'explanation': None,
                        'topic': detect_network_topic(current_question)
                    })
        
        # Handle tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            
            # Add table content to the last question's answer if exists
            if questions and table_text:
                questions[-1]['answer'] += '\n\n' + '\n'.join(table_text)
    
    except Exception as e:
        print(f"Error parsing {filepath}: {e}")
    
    return questions

def parse_basic_networking_file(filepath, reference_questions, seen_questions=None):
    """Parse Basic_Networking_Questions.docx and match with reference answers"""
    questions = []
    if seen_questions is None:
        seen_questions = set()
    
    try:
        doc = Document(filepath)
        current_topic = 'Networking Basics'
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this is a category heading (bold, no number)
            is_bold = False
            if para.runs:
                is_bold = any(run.bold for run in para.runs)
            
            if is_bold and not re.match(r'^\d+\.', text) and len(text) < 100:
                # This is a category heading
                current_topic = text
                continue
            
            # Check if this is a question (starts with bullet, number, or question word)
            if re.match(r'^[\-\•\*]', text) or re.match(r'^\d+[\.\)]', text) or text.lower().startswith(('what', 'how', 'why', 'explain', 'define', 'describe')):
                # Clean up the question
                question = re.sub(r'^[\-\•\*\d\.\)]+\s*', '', text)
                
                # Check for duplicates using hash
                q_hash = hash(question.lower().strip())
                if q_hash in seen_questions:
                    continue  # Skip duplicate
                
                # Try to find matching answer from reference questions
                answer = 'Refer to networking study material'
                best_match_score = 0
                
                for ref_q in reference_questions:
                    score = similarity(question, ref_q['question'])
                    if score > best_match_score and score > 0.6:
                        best_match_score = score
                        answer = ref_q['answer']
                
                if len(question) > 10:
                    seen_questions.add(q_hash)
                    questions.append({
                        'question_type': 'theory',
                        'question': question,
                        'option_a': None,
                        'option_b': None,
                        'option_c': None,
                        'option_d': None,
                        'answer': answer,
                        'explanation': None,
                        'topic': current_topic
                    })
    
    except Exception as e:
        print(f"Error parsing {filepath}: {e}")
    
    return questions

def parse_network_files():
    """Parse both networking files and return combined questions"""
    questions = []
    seen_questions = set()  # Global duplicate tracking
    
    # Parse Fresher_Networking_Interview.docx first (has answers)
    fresher_file = os.path.join('data', 'Fresher Networking Interview.docx')
    if os.path.exists(fresher_file):
        fresher_questions = parse_fresher_networking_file(fresher_file, seen_questions)
        questions.extend(fresher_questions)
        print(f"Parsed {len(fresher_questions)} questions from Fresher Networking Interview.docx")
        
        # Parse Basic_Networking_Questions.docx and match with fresher questions
        basic_file = os.path.join('data', 'Basic Networking Questions.docx')
        if os.path.exists(basic_file):
            basic_questions = parse_basic_networking_file(basic_file, fresher_questions, seen_questions)
            questions.extend(basic_questions)
            print(f"Parsed {len(basic_questions)} questions from Basic Networking Questions.docx")
    
    return questions
